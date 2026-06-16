# -*- coding: utf-8 -*-
"""Manual detallado de pruebas AZUR ERP en Word (.docx), brandeado, español Perú."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

AZUR = RGBColor(0xE2, 0x06, 0x27)
AZUR2 = RGBColor(0xBE, 0x17, 0x23)
GRIS = RGBColor(0x66, 0x66, 0x66)
VERDE = RGBColor(0x0A, 0x7D, 0x33)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)
LOGO = "logoazur.png"

doc = Document()
st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10.5)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexcolor); tcPr.append(shd)

def cell_text(cell, text, bold=False, color=None, size=9.5, align=None, white=False):
    cell.text = ""; p = cell.paragraphs[0]
    if align: p.alignment = align
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
    if white: r.font.color.rgb = BLANCO
    elif color: r.font.color.rgb = color

def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs: r.font.color.rgb = AZUR; r.font.size = Pt(16)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs: r.font.color.rgb = AZUR2; r.font.size = Pt(13)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    for r in p.runs: r.font.color.rgb = RGBColor(0x33,0x33,0x33); r.font.size = Pt(11.5)
    return p

def para(text, color=GRIS, size=10):
    p = doc.add_paragraph(); r = p.add_run(text); r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return p

def proposito(text):
    p = doc.add_paragraph(); r = p.add_run("Para qué sirve: "); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = AZUR2
    r2 = p.add_run(text); r2.font.size = Pt(10)
    return p

def tabla(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    for j, x in enumerate(headers):
        c = t.rows[0].cells[j]; shade(c, "BE1723"); cell_text(c, x, bold=True, white=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        rc = t.add_row()
        for j, val in enumerate(row):
            cell_text(rc.cells[j], val, size=9.5)
    if widths:
        for col, w in zip(t.columns, widths):
            for c in col.cells: c.width = w
    doc.add_paragraph()
    return t

def campos(rows):
    h3("Campos del formulario")
    tabla(["Campo", "Para qué sirve / qué escribir", "Formato · obligatorio"], rows,
          [Inches(1.5), Inches(3.3), Inches(1.6)])

def botones(rows):
    h3("Botones y acciones")
    tabla(["Botón / acción", "Qué hace al pulsarlo"], rows, [Inches(2.0), Inches(4.5)])

def pestanas(rows):
    h3("Pestañas / secciones")
    tabla(["Pestaña", "Qué contiene"], rows, [Inches(2.0), Inches(4.5)])

def pasos(items):
    h3("Paso a paso")
    for accion, esperado in items:
        p = doc.add_paragraph(style="List Number"); r = p.add_run(accion); r.font.size = Pt(10.5)
        if esperado:
            sp = doc.add_paragraph(); sp.paragraph_format.left_indent = Inches(0.5)
            rr = sp.add_run("✓ Debes ver: "); rr.bold = True; rr.font.size = Pt(9); rr.font.color.rgb = VERDE
            r2 = sp.add_run(esperado); r2.font.size = Pt(9); r2.font.color.rgb = GRIS

def obs(filas=3):
    p = doc.add_paragraph(); r = p.add_run("📝 Observaciones del cliente sobre esta sección:"); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = AZUR2
    t = doc.add_table(rows=1, cols=4); t.style = "Table Grid"
    for j, x in enumerate(["Qué probaste", "¿Funcionó? Sí/No", "Severidad A/M/B", "Comentario"]):
        c = t.rows[0].cells[j]; shade(c, "E20627"); cell_text(c, x, bold=True, white=True, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(filas):
        rc = t.add_row()
        for j in range(4): cell_text(rc.cells[j], "", size=10)
    for col, w in zip(t.columns, [Inches(1.9), Inches(1.1), Inches(1.2), Inches(2.3)]):
        for c in col.cells: c.width = w
    doc.add_paragraph()

def nota(text):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run("💡 "); r.font.size = Pt(9.5)
    r2 = p.add_run(text); r2.font.size = Pt(9.5); r2.italic = True; r2.font.color.rgb = RGBColor(0x33,0x55,0x88)

def regla(text):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run("⚠ Regla/validación: "); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = AZUR
    r2 = p.add_run(text); r2.font.size = Pt(9.5)

def ejemplo(titulo, headers, rows):
    p = doc.add_paragraph(); r = p.add_run("📐 Ejemplo: " + titulo); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = VERDE
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    for j, x in enumerate(headers):
        c = t.rows[0].cells[j]; shade(c, "0A7D33"); cell_text(c, x, bold=True, white=True, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        rc = t.add_row()
        for j, val in enumerate(row): cell_text(rc.cells[j], val, size=9, align=(WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.RIGHT))
    doc.add_paragraph()

def casos(items):
    p = doc.add_paragraph(); r = p.add_run("✅ Casos de prueba (hazlos y verifica el resultado):"); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = AZUR2
    for txt, esperado in items:
        b = doc.add_paragraph(style="List Bullet"); rb = b.add_run(txt); rb.font.size = Pt(10)
        sp = doc.add_paragraph(); sp.paragraph_format.left_indent = Inches(0.5)
        rr = sp.add_run("→ Resultado esperado: "); rr.bold = True; rr.font.size = Pt(9); rr.font.color.rgb = VERDE
        r2 = sp.add_run(esperado); r2.font.size = Pt(9); r2.font.color.rgb = GRIS

def add_toc():
    p = doc.add_paragraph(); run = p.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
    tt = OxmlElement("w:t"); tt.text = "Abre el documento en Word y presiona F9 (o clic derecho → Actualizar campos) para ver el índice con páginas."
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    for e in (f1, instr, f2, tt, f3): run._r.append(e)

def page_number_footer():
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("AZUR Constructora e Inmobiliaria · Manual de Pruebas del ERP · Pág. ")
    r.font.size = Pt(8); r.font.color.rgb = GRIS
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    fp._p.append(fld)

# ───────────────────────── PORTADA ─────────────────────────
if os.path.exists(LOGO):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(LOGO, width=Inches(2.2))
doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Manual de Pruebas — Plataforma ERP + App de Obra"); r.bold = True; r.font.size = Pt(23); r.font.color.rgb = AZUR
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("AZUR Constructora e Inmobiliaria"); r.font.size = Pt(14); r.font.color.rgb = AZUR2
d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = d.add_run("Guía detallada, campo por campo y botón por botón, para probar TODOS los flujos\nde extremo a extremo y registrar observaciones en este mismo documento."); r.font.size = Pt(11); r.font.color.rgb = GRIS
doc.add_paragraph()
t = doc.add_table(rows=3, cols=2); t.style = "Table Grid"
for i, (k, v) in enumerate([("Empresa / cliente", ""), ("Persona que prueba", ""), ("Fecha de la prueba", "")]):
    cell_text(t.rows[i].cells[0], k, bold=True, size=10); cell_text(t.rows[i].cells[1], v, size=10)
doc.add_paragraph()
para("Contraseña de todos los usuarios de prueba: Azur2026!", color=AZUR2, size=10)
doc.add_page_break()
h1("Índice"); add_toc(); doc.add_page_break()

# ───────────────────── 1. CONCEPTOS ─────────────────────
h1("1. Conceptos básicos del sistema")
proposito("Entender cómo está organizado el negocio dentro del sistema antes de probar.")
para("La jerarquía del sistema es: Línea de negocio → Proyecto → Etapas/Partidas → (Solicitudes de pago · Valorizaciones · Partes diarios).")
h3("Líneas de negocio")
tabla(["Línea", "Para qué"], [
    ["Azur Construcción", "Obras con presupuesto por partidas, APU, curva S y valorización (proyecto Grande)."],
    ["Cocina Pro", "Línea propia con su catálogo e identidad."],
    ["Mantenimiento", "Servicios recurrentes con cronograma de visitas (proyecto Chico)."],
])
h3("Tipos de proyecto")
tabla(["Tipo", "Cómo se comporta"], [
    ["Grande (construcción)", "Usa metrados, APU, curva S, valorización formal, adicionales y liquidación."],
    ["Chico (mantenimiento/acondicionamiento)", "Lista de servicios, control de caja/gasto; activa la pestaña Mantenimiento."],
])
h3("Roles y qué hace cada uno")
tabla(["Rol", "Qué puede hacer"], [
    ["Gerencia General", "Ve todo, aprobación final de pagos, dashboards, márgenes."],
    ["Jefe de Proyectos", "Proyectos, valorizaciones, 1ª aprobación de pagos."],
    ["Jefe de Presupuestos", "Cotizaciones (APU) y catálogos."],
    ["Administrador (Pamela)", "Finanzas: programa y paga, cajas, CxC y CxP."],
    ["Comercial", "Cotizaciones y su estado."],
    ["Residente / Coordinador", "App de obra: check-in, RDO, evidencias, solicitudes, tareo."],
    ["Prevencionista (SOMA)", "Seguridad: charlas, observaciones, incidentes."],
    ["Logístico", "Almacén y compras."],
])
obs(2)

# ───────────────────── 2. ACCESO ─────────────────────
h1("2. Acceso y navegación")
proposito("Ingresar al sistema y reconocer el menú.")
pasos([
    ("Abre la dirección web de la plataforma en una computadora (y en el celular para la app de obra).", "Pantalla de inicio de sesión con el logo de AZUR sobre fondo blanco y, debajo, la lista de usuarios de prueba."),
    ("Toca un usuario de prueba (o escribe correo y contraseña Azur2026!) y pulsa Ingresar.", "Entras: los roles de oficina ven el menú lateral (sidebar); los de obra ven la app de campo con barra inferior."),
    ("Observa arriba a la derecha la campana de notificaciones y tu nombre/menú de usuario.", "La campana muestra un número si hay notificaciones sin leer; el menú permite ir a Mi perfil o Cerrar sesión."),
    ("Prueba Cerrar sesión desde el menú de usuario.", "Vuelves a la pantalla de inicio de sesión."),
])
botones([
    ["Menú lateral (Dashboard, Comercial, Proyectos, Finanzas, Almacén, Reportes, Alertas, Clientes, Catálogos, Usuarios)", "Cada opción abre su módulo. Solo aparecen las que tu rol puede usar."],
    ["Campana 🔔", "Abre el panel de notificaciones; al tocar una, se marca leída."],
    ["Menú de usuario (tu nombre)", "Mi perfil / Cerrar sesión."],
])
obs(2)

# ───────────────────── 3. DASHBOARD ─────────────────────
h1("3. Dashboard (inicio)")
proposito("Ver la película del negocio: ingresos/egresos del mes, salud de cada proyecto y alertas.")
para("Entra como Gerencia y abre Dashboard (primera opción del menú).")
h3("Qué muestra")
tabla(["Elemento", "Qué significa"], [
    ["Ingresos del mes", "Suma de abonos del cliente del mes en curso."],
    ["Egresos del mes", "Suma de pagos realizados en el mes."],
    ["Proyectos activos / total", "Cuántos están en ejecución."],
    ["Alertas abiertas", "Alertas críticas sin resolver."],
    ["Barra de 3 tramos (por proyecto)", "Proyectado (referencia) · Pagos/abonos · Gasto. El color indica salud (verde/ámbar/rojo)."],
    ["Panel de alertas críticas", "Las últimas alertas del sistema."],
])
para("Regla de salud #1: el gasto no debe superar lo cobrado, y lo cobrado debe avanzar hacia lo proyectado. Regla #2: el gasto real no debe superar lo valorizado. Si se rompen, la barra se pone roja y se genera alerta.")
pasos([
    ("Pulsa una barra de proyecto.", "Te lleva al detalle de ese proyecto."),
])
obs(2)

# ───────────────────── 4. CLIENTES ─────────────────────
h1("4. Maestros — Clientes")
proposito("Registrar y administrar la cartera de clientes (se usan al cotizar).")
para("Menú Clientes (grupo Maestros). Como Comercial, Gerencia, Presupuestos o Administrador.")
campos([
    ["Razón social", "Nombre o razón social del cliente.", "Texto · obligatorio"],
    ["Tipo doc", "RUC, DNI o CE.", "Lista"],
    ["RUC / DNI", "Número de documento.", "Solo números (RUC 11, DNI 8)"],
    ["Contacto", "Persona de contacto.", "Texto"],
    ["Teléfono", "Teléfono del contacto.", "Solo números"],
    ["Email", "Correo del contacto.", "Formato correo válido"],
    ["Origen", "Cómo llegó el cliente (directo, recomendación, oficina, llamada).", "Lista"],
    ["Ubicación (referencial)", "Dirección o referencia en texto libre.", "Texto"],
])
botones([
    ["Nuevo cliente", "Abre el formulario para crear un cliente."],
    ["Importar", "Pega varias filas (Razón social, RUC, contacto) para cargar la cartera en lote; detecta duplicados por RUC."],
    ["Buscar", "Filtra la lista por razón social, RUC o contacto."],
    ["Lápiz (editar)", "Abre el cliente para modificarlo."],
])
pasos([
    ("Pulsa Nuevo cliente, llena los campos y Guardar.", "El cliente aparece en la tabla; al escribir letras en RUC/teléfono no las acepta."),
    ("Usa el buscador con parte del nombre.", "La lista se filtra al instante."),
    ("Pulsa Importar, pega 2-3 líneas separadas por tabulación y Importar.", "Muestra cuántos importó y cuántos duplicados omitió."),
    ("Edita un cliente con el lápiz y guarda un cambio.", "El cambio queda guardado."),
])
obs(3)

# ───────────────────── 5. CATÁLOGOS ─────────────────────
h1("5. Maestros — Catálogos")
proposito("Mantener partidas, insumos, contratistas/proveedores, plantillas y medios de pago.")
para("Menú Catálogos (grupo Maestros). Tiene varias pestañas.")
pestanas([
    ["Clientes", "Misma cartera de clientes (también editable aquí)."],
    ["Contratistas/Proveedores", "Razón social, tipo, RUC/DNI, especialidad, contacto, teléfono, banco, cuenta, CCI."],
    ["Partidas", "Línea, código, descripción, unidad y costo referencial. Base para cotizar."],
    ["Insumos", "Código, nombre, unidad, precio y tipo (material/mano de obra/equipos)."],
    ["Plantillas", "Textos de condiciones, servicios incluidos/omitidos y garantía precargados."],
    ["Medios de pago", "Cuentas de la empresa (banco, titular, cuenta soles/dólares, CCI, detracción)."],
])
botones([
    ["Nuevo (en cada pestaña)", "Crea un registro en ese catálogo."],
    ["Lápiz / Tacho", "Editar o eliminar un registro."],
    ["Actualizar precios", "Ajusta en % (ej. +5%) los precios de insumos/partidas y avisa cuántas cotizaciones podrían quedar desactualizadas."],
])
pasos([
    ("En Partidas crea una partida nueva (línea, código, descripción, unidad, costo).", "Queda guardada y aparecerá al cotizar."),
    ("Pulsa Actualizar precios, elige Insumos, escribe 5 y Aplicar.", "Sube 5% los precios y muestra un aviso."),
])
obs(3)

# ───────────────────── 6. COMERCIAL ─────────────────────
h1("6. Comercial — Cotizaciones (detallado)")
proposito("Es el punto de entrada del sistema. Aquí se arma el presupuesto de venta; al aprobarse, la cotización se convierte automáticamente en un proyecto.")
para("La cotización se construye como un ÁRBOL de hasta 4 niveles (Partida general → Sub-partida → Actividad → Sub-actividad). La pantalla tiene, lado a lado, el CUADRO DE COSTOS (lo que cuesta hacerlo) y el CUADRO DE MARGEN (el precio que ve el cliente). Regla de oro: solo el ÚLTIMO nivel de cada rama (la 'hoja') captura unidad, cantidad, costo unitario y % de margen; los niveles superiores se BLOQUEAN y se calculan sumando a sus hijos.")

h2("6.1 Listado de cotizaciones")
para("Tabla con código, proyecto, cliente, línea, estado (con color) y fecha.")
botones([
    ["Nueva cotización", "Abre el asistente de creación (datos generales)."],
    ["Buscar", "Filtra en el servidor por nombre de proyecto, código o asunto."],
    ["Menú ⋮ de fila", "Abrir/Editar la cotización, o Eliminar (solo si NO fue aceptada)."],
    ["Anterior / Siguiente", "Paginación (20 por página) cuando hay muchas."],
])
casos([
    ("Escribe en el buscador parte del nombre de un proyecto.", "La lista se filtra y la paginación se reinicia."),
    ("Abre el menú ⋮ de una cotización aceptada.", "La opción Eliminar NO aparece (ya generó proyecto)."),
])
obs(2)

h2("6.2 Nueva cotización — datos generales (Paso 0 y 1)")
campos([
    ["Origen del lead", "Cómo llegó el contacto. Alimenta la reportería comercial.", "Lista (directo/recomendación/oficina/llamada) · obligatorio"],
    ["Cliente", "Cliente al que se cotiza. Con el botón + creas uno sin salir.", "Lista · obligatorio"],
    ["Línea de negocio", "Define plantilla, catálogo e identidad (Azur Construcción / Cocina Pro / Mantenimiento).", "Lista · obligatorio"],
    ["Tipo de cotización", "Única de obra = proyecto Grande; Programada/Recurrencia = proyecto Chico (mantenimiento).", "Lista · obligatorio"],
    ["Nombre del proyecto", "Cómo se identificará el trabajo.", "Texto · obligatorio"],
    ["Asunto", "Breve descripción (ej. 'instalación de viniles').", "Texto"],
    ["Ubicación + Mapa", "Busca la dirección (tipo Google) o haz clic en el mapa; guarda la coordenada y rellena la ubicación.", "Texto + mapa"],
    ["Vigencia (días)", "Días que la cotización es válida (dispara el aviso de 'por vencer').", "Número"],
    ["Plantilla de condiciones", "Precarga condiciones/servicios/garantía (editables) o 'página en blanco'.", "Lista"],
])
casos([
    ("Llena los datos, ubica en el mapa y pulsa 'Crear y armar presupuesto'.", "Se crea con código correlativo COT-#### y abre el editor."),
    ("Pulsa + junto a Cliente y crea uno nuevo (razón social + RUC de 11 dígitos).", "Queda creado y seleccionado; el RUC solo acepta números."),
])
obs(2)

h2("6.3 Editor — árbol de costos y margen (el corazón)")
h3("Qué significa cada columna")
tabla(["Columna", "Qué es y de dónde sale"], [
    ["Ítem", "Numeración automática según el nivel: 1.0 (partida general), 1.1 (sub-partida), 1.1.1 (actividad), 1.1.1.1 (sub-actividad)."],
    ["Título", "Nombre editable de la fila."],
    ["Und / Cant / C. Unit", "Unidad, cantidad y costo unitario. SOLO se escriben en la hoja (último nivel)."],
    ["Subtotal (costos)", "= Cantidad × Costo unitario. En niveles superiores = suma de hijos."],
    ["% Marg", "Margen por partida (editable solo en la hoja). Mínimo sugerido 30%."],
    ["P. Unit (con margen)", "Precio unitario = Costo unitario ÷ (1 − %margen)."],
    ["Precio (subtotal c/margen)", "= P. Unit × Cantidad. Es lo único que ve el cliente."],
])
regla("Solo la hoja captura datos. Si una partida tiene hijos, sus celdas de unidad/cantidad/C.U./margen se BLOQUEAN y su valor es la suma de los hijos. El margen NUNCA es un costo: es la diferencia entre precio y costo.")
para("Fórmulas exactas que usa el sistema:")
tabla(["Concepto", "Fórmula"], [
    ["Precio unitario", "C.U. ÷ (1 − %margen)"],
    ["Margen (monto)", "(Precio unitario − C.U.) × Cantidad"],
    ["Precio (subtotal con margen)", "Precio unitario × Cantidad"],
    ["Total de un nivel superior", "Suma de los precios de sus hijos"],
])
ejemplo("Actividad de concreto con 10% de margen", ["Dato", "Valor"], [
    ["Cantidad", "30 m³"], ["Costo unitario", "S/ 4,500.00"], ["Subtotal de costo (30×4500)", "S/ 135,000.00"],
    ["% Margen", "10%"], ["Precio unitario (4500 ÷ 0.9)", "S/ 5,000.00"], ["Margen monto ((5000−4500)×30)", "S/ 15,000.00"],
    ["Precio con margen (5000×30)", "S/ 150,000.00"],
])
botones([
    ["Agregar partida", "Abre el selector de catálogo para crear una partida de nivel 1."],
    ["+ (icono en la fila)", "Agrega un hijo del nivel siguiente (hasta el nivel 4)."],
    ["Capas (Detallar APU)", "Abre el desglose del costo unitario (ver 6.4). Si la partida tiene APU, el C.U. queda en rojo y de solo lectura."],
    ["Tacho", "Elimina la fila y todos sus hijos."],
    ["Selector de catálogo", "Buscador por código/descripción; cada ítem muestra su precio y un distintivo 'APU' si trae desglose. También 'Crear en blanco'."],
])
casos([
    ("Pulsa 'Agregar partida' y elige del catálogo una con distintivo APU (ej. Concreto f'c=210).", "Se agrega con su C.U. ya calculado desde el APU y el icono de capas queda en rojo."),
    ("Con el + crea una sub-partida, luego una actividad y una sub-actividad.", "La numeración se arma sola: 1.1, 1.1.1, 1.1.1.1; no deja pasar del nivel 4."),
    ("En una hoja escribe Cantidad 30, C.U. 4500 y % margen 10.", "Subtotal 135,000; P.Unit 5,000; Precio 150,000 al instante."),
    ("Intenta escribir cantidad/C.U. en una partida que tiene hijos.", "No se puede: esos campos están bloqueados y muestran la suma de los hijos."),
])
obs(3)

h2("6.4 Detallar APU (Análisis de Precios Unitarios)")
proposito("Construir el costo unitario de una partida a partir de sus componentes (mano de obra, materiales, equipos, subcontratos, gastos generales). El C.U. de la partida pasa a ser la suma de los componentes.")
campos([
    ["Tipo", "Categoría del componente.", "Lista (mano de obra/materiales/equipos/subcontratos/GG)"],
    ["Descripción", "Insumo o recurso (ej. 'Cemento Portland').", "Texto · obligatorio"],
    ["Unidad", "Unidad del componente (bls, m³, hh…).", "Texto"],
    ["Cantidad/und", "Cuánto de ese insumo entra por UNA unidad de la partida.", "Número"],
    ["Precio", "Precio unitario del insumo.", "Número"],
    ["Cuadrilla / Rendimiento", "Opcionales, de referencia para mano de obra.", "Número"],
])
ejemplo("APU de 1 m³ de concreto (C.U. resultante = suma de parciales)", ["Componente", "Cant × Precio = Parcial"], [
    ["Cemento (materiales)", "9 bls × 32.5 = 292.50"], ["Arena gruesa", "0.5 m³ × 65 = 32.50"],
    ["Piedra chancada", "0.55 m³ × 75 = 41.25"], ["Cuadrilla concreto (MO)", "2.5 hh × 20 = 50.00"],
    ["Mezcladora (equipo)", "0.4 hm × 35 = 14.00"], ["C.U. de la partida", "= 430.25"],
])
botones([
    ["Agregar", "Suma el componente; el C.U. de la partida se recalcula automáticamente."],
    ["Guardar como plantilla", "Crea una partida en el catálogo maestro con este APU para reutilizarla en otras cotizaciones/proyectos."],
    ["X", "Elimina el componente."],
])
casos([
    ("Agrega 3-4 componentes y observa el C.U.", "El C.U. de la partida es la suma de los parciales y queda de solo lectura."),
    ("Pulsa 'Guardar como plantilla' y ponle un código.", "Aparece como nueva partida en Catálogos con el distintivo APU."),
])
obs(2)

h2("6.5 Bloque de totales, márgenes y descuento")
para("Debajo del árbol, el sistema arma la secuencia de totales. GG, GA, Utilidad e IGV son porcentajes configurables y se pueden OCULTAR al cliente con interruptores (algunos clientes no quieren ver utilidad/GG; igual se cobran dentro del precio).")
ejemplo("Secuencia de totales (números del modelo real ADECCO)", ["Concepto", "Importe S/"], [
    ["Subtotal (con margen)", "3,354,299.30"], ["Gastos generales (5%)", "167,714.96"],
    ["Gastos administrativos (5%)", "167,714.96"], ["Utilidad (5%)", "167,714.96"],
    ["Costo directo", "3,857,444.19"], ["I.G.V. (18%)", "694,339.95"],
    ["TOTAL", "4,551,784.15"], ["Descuento comercial (5%)", "227,589.21"],
    ["TOTAL CON DESCUENTO", "4,324,194.94"],
])
regla("El descuento comercial NO está visible por defecto (estrategia). Se activa solo cuando el cliente pide rebaja, y se absorbe del MARGEN: el costo directo interno nunca cambia. La 'vista interna' muestra el costo directo real y el margen total, que el cliente jamás ve.")
botones([
    ["Mostrar GG / GA / Utilidad / IGV (interruptores)", "Incluyen u ocultan cada concepto en el PDF del cliente."],
    ["Agregar descuento comercial", "Activa un % de descuento sobre el total (aparece TOTAL CON DESCUENTO)."],
])
casos([
    ("Apaga el interruptor de Utilidad y de GG.", "Desaparecen de la lista y del PDF, pero el TOTAL no cambia."),
    ("Activa un descuento de 5%.", "Aparece 'Total con descuento'; el costo directo real (vista interna) se mantiene igual."),
])
obs(2)

h2("6.6 Condiciones, forma de pago y medios de pago")
campos([
    ["Forma de pago (filas)", "Concepto y % de cada pago (ej. 20% adelanto + 80% valorizaciones; o 4 armadas de 25%).", "Texto + % · la suma no puede pasar de 100%"],
    ["Plazo de ejecución + tipo", "N° de días y si son calendario o útiles.", "Número + lista"],
    ["Condiciones generales", "Texto precargado de la plantilla, editable.", "Texto largo"],
    ["Servicios incluidos / omitidos", "Qué cubre y qué no el servicio.", "Texto largo"],
    ["Garantía", "Texto de garantía con interruptor para incluirla o no en el PDF.", "Texto + sí/no"],
    ["Medios de pago", "Cuentas de la empresa (Interbank, BBVA, detracción del Banco de la Nación).", "Solo lectura"],
])
regla("La suma de los % de la forma de pago se valida: si pasa de 100% no deja guardar.")
casos([
    ("Agrega pagos que sumen 110% e intenta guardar.", "El sistema advierte que supera el 100% y no guarda."),
    ("Edita 'Servicios incluidos' y apaga la garantía.", "Se guarda solo (autosave); la garantía no saldrá en el PDF."),
])
obs(2)

h2("6.7 Estados de la cotización")
tabla(["Estado", "Qué significa / qué lo activa"], [
    ["Borrador", "Edición libre, sin alertas."],
    ["Enviada", "Se mandó el PDF al cliente; corre el cronómetro de vigencia."],
    ["En negociación", "El cliente pide ajustes; aquí se activa el descuento y se guardan versiones."],
    ["Aceptada", "Se convierte en proyecto (sin margen) + cronograma de cobros + caja chica."],
    ["Vencida", "Pasó la vigencia sin respuesta (lo marca el proceso diario automático)."],
    ["Rechazada", "El cliente declinó; se registra el motivo."],
])
botones([
    ["Generar PDF (cliente)", "PDF brandeado con SOLO las columnas del cliente (ítem, descripción, unidad, cantidad, P.U., subtotal y totales). Nunca muestra costo ni margen."],
    ["Descargar Excel (interno)", "Excel brandeado con la matriz completa: costos, % margen y márgenes."],
    ["Enviar por WhatsApp", "Abre WhatsApp con un mensaje pre-formateado y el enlace."],
    ["Guardar versión", "Guarda una foto de la negociación (v1, v2, v3…) con justificación."],
    ["Marcar como enviada / Pasar a negociación", "Cambian el estado."],
    ["Aprobar → crear proyecto", "Crea el proyecto con el itemizado a COSTO (sin margen), el cronograma de cobros (armadas) y la caja chica; avisa a Jefe de Proyectos y Presupuestos."],
    ["Rechazar / Eliminar cotización", "Rechaza con motivo / elimina (solo si no fue aceptada)."],
])
h3("Edición colaborativa en tiempo real")
para("Varios miembros del equipo (Comercial/Presupuestos/Gerencia) pueden abrir la MISMA cotización a la vez. Arriba se ven los avatares de quién está conectado y los cambios de uno aparecen en las pantallas de los demás en segundos, sin recargar (autosave).")
h3("Historial de modificaciones")
para("Pestaña aparte que registra, por cada cambio: usuario, fecha/hora, la partida afectada y el campo con su valor anterior → nuevo. Cada usuario tiene un color. Permite FILTRAR por usuario y REVERTIR un cambio (restaura el valor anterior; queda auditado).")
casos([
    ("Abre la cotización en dos ventanas con usuarios distintos y edita una fila en una.", "En la otra ventana ves el avatar del otro usuario y el cambio aparece en segundos."),
    ("Cambia la cantidad de una partida, ve a Historial, filtra por tu usuario y pulsa Revertir.", "El historial muestra 'Partida X · Cantidad: 1 → 2' y al revertir vuelve al valor anterior."),
    ("Genera el PDF (cliente) y ábrelo.", "NO aparece ninguna columna de costo ni de margen; solo precios al cliente y el logo de AZUR."),
    ("Pulsa 'Aprobar → crear proyecto' y revisa el proyecto creado.", "El itemizado llegó a COSTO (sin margen), con cronograma de cobros y caja chica; llegó notificación al Jefe."),
])
obs(3)

# ───────────────────── 7. PROYECTOS ─────────────────────
h1("7. Proyectos")
proposito("Gobernar la ejecución de la obra: presupuesto, programación, valorización, finanzas y cierre.")
h2("7.1 Listado")
botones([["Buscar / Paginación", "Filtra por nombre o código; navega por páginas."], ["Clic en el código", "Abre el proyecto."]])

h2("7.2 Detalle — pestañas")
pestanas([
    ["Resumen", "KPIs (contrato, cobrado, gasto, caja), barra de 3 tramos, Curva S, hitos, informe de obra y configuración rápida."],
    ["Last Planner", "Itemizado + programación + valorizaciones + estado/prioridad automáticos."],
    ["Cronograma de cobros", "Armadas (% por avance o fecha) con su monto."],
    ["Adicionales", "Adicionales/deductivos con aprobación."],
    ["Equipo", "Personas asignadas con su rol de obra."],
    ["Mantenimiento", "(Solo proyectos chicos) cronograma de servicios recurrentes."],
    ["Campo", "Lo capturado desde la app: asistencia (GPS), partes diarios, evidencias y SST."],
    ["Liquidación", "Balance final del proyecto y saldo del adelanto."],
    ["Expediente", "Documentos por carpetas (contratos, planos, etc.)."],
])
obs(2)

h2("7.3 Resumen")
botones([
    ["Informe de obra (PDF)", "Abre un cuadro para elegir qué secciones incluir y genera el PDF para el cliente."],
    ["Agregar hito", "Registra un hito con fecha; se marca Próximo/Vencido/Cumplido."],
    ["Configuración rápida", "Cambia estado, contrato total, % de adelanto y tope de caja."],
])
obs(2)

h2("7.4 Last Planner — los 4 cuadrantes (detallado)")
para("Es una matriz horizontal con 4 bloques que se leen de izquierda a derecha. Mismo árbol de 4 niveles que la cotización, pero a COSTO (sin margen). La jefatura llena los cuadrantes 1 y 2; el residente/jefe registra el 3; el 4 es automático.")
h3("Cuadrante 1 — Presupuesto / itemizado")
tabla(["Columna", "Qué es"], [
    ["Ítem / Título", "Numeración automática (1.0/1.1/1.1.1/1.1.1.1) y nombre."],
    ["Und / Cant / C. Unit", "Unidad, cantidad y costo (solo en la hoja; editable)."],
    ["Subtotal", "Monto de la fila (en sub-partidas/actividades)."],
    ["Total", "Se llena solo en las PARTIDAS GENERALES (suma de toda su rama)."],
])
regla("Puedes traer partidas del catálogo (con APU) o crearlas; el itemizado del proyecto puede diferir del comercial (Proyectos arma su propia estructura). También se puede detallar APU por partida con el icono de capas.")
h3("Cuadrante 2 — Programación")
tabla(["Columna", "Qué es"], [
    ["Contratista responsable", "Se elige de la lista de contratistas/proveedores registrados."],
    ["Fecha de inicio / entrega", "Definen el periodo de la partida."],
    ["Duración (Dur)", "Días de la partida; junto con el inicio define el avance proyectado por semana."],
])
h3("Cuadrante 3 — Valorizaciones semanales acumulables")
para("Cada semana se agrega una valorización (dos columnas: % avance y Total). Solo se ingresa el % de avance de la semana en la HOJA; el sistema arrastra (suma) hacia arriba.")
tabla(["Cálculo", "Fórmula"], [
    ["Total de la semana (por hoja)", "% de la semana × TOTAL de la partida"],
    ["% acumulado de la partida", "Suma de los % de todas sus semanas"],
    ["Valorizado acumulado", "% acumulado × TOTAL de la partida"],
    ["Saldo", "TOTAL × (1 − % acumulado)"],
    ["Rollup a niveles superiores", "Ponderado por monto: %acum(padre)= Σ(Total_hijo×%acum_hijo) ÷ Σ(Total_hijo)"],
])
ejemplo("Sub-partida de S/ 240,000 valorizada en 3 semanas (25% / 25% / 50%)", ["Semana", "% / Total / %acum / Saldo"], [
    ["Semana 1 (25%)", "S/ 60,000 · acum 25% · saldo 180,000"],
    ["Semana 2 (25%)", "S/ 60,000 · acum 50% · saldo 120,000"],
    ["Semana 3 (50%)", "S/ 120,000 · acum 100% · saldo 0"],
])
regla("El % acumulado de una partida nunca puede pasar de 100%: si la suma de semanas lo supera, el sistema lo bloquea. El saldo nunca es negativo.")
h3("Cuadrante 4 — Estado y Prioridad (AUTOMÁTICOS)")
para("No se escriben a mano: el sistema los calcula comparando el avance REAL con el PROYECTADO. El proyectado se reparte proporcional entre las semanas de duración (ej. 5 semanas → 20% por semana: 20/40/60/80/100% acumulado).")
tabla(["Estado", "Cuándo se asigna"], [
    ["Completado", "% acumulado = 100%."],
    ["En progreso", "Avanzó algún % entre una valorización y la siguiente."],
    ["Detenido", "El % no avanzó respecto a la valorización anterior."],
    ["En espera", "La fecha de inicio aún no llega."],
    ["Pendiente", "No tiene fecha de inicio."],
    ["Retrasado", "Venció el plazo y no llegó al 100%."],
    ["Cancelado", "Override manual."],
])
tabla(["Prioridad (real ÷ proyectado de la semana)", "Resultado"], [
    ["Menos del 50% de lo proyectado", "Muy alta"],
    ["Entre 50% y por debajo del objetivo", "Alta"],
    ["Justo en lo proyectado (≈100%)", "Media"],
    ["Por encima del objetivo", "Baja"],
    ["Muy por encima (≥120%)", "Muy baja"],
    ["Sin fechas / proyectado 0", "Media (por defecto)"],
])
ejemplo("Caso del modelo (proyectado 20/40/60/80/100; real 0.3/0.4/0.4/0.4/0.9/1.0)", ["Semana", "Prioridad → Estado"], [
    ["S1 real 30% vs 20%", "Muy baja → En progreso"],
    ["S2 real 40% vs 40%", "Media → En progreso"],
    ["S3 real 40% vs 60%", "Alta → Detenido (no avanzó)"],
    ["S4 real 40% vs 80%", "Muy alta → Detenido"],
    ["S5 real 90% vs 100%", "Alta → Retrasado"],
    ["S6 real 100%", "Media → Completado"],
])
h3("% Acumulado y Saldo")
para("La columna % Acum muestra una barra verde proporcional al avance (como en el Excel) y el Saldo es lo que falta por valorizar.")
h3("Botones")
botones([
    ["Partida", "Agrega una partida general desde el catálogo (o en blanco)."],
    ["+ (en fila)", "Agrega un hijo (hasta nivel 4)."],
    ["Capas (APU)", "Detalla el APU de esa partida del proyecto."],
    ["Tacho", "Elimina la fila."],
    ["Nueva valorización", "Crea las columnas (% y Total) de la siguiente semana."],
    ["Guardar avances", "Guarda los % de la semana; recalcula %acum, saldo, estado y prioridad, y genera alertas de salud si corresponde."],
    ["Resumen PDF", "Genera el PDF/resumen ejecutivo de la valorización para el cliente."],
    ["Registrar cobro", "Lleva el cobro neto de la valorización a la caja del proyecto."],
])
h3("Dilución del adelanto (al registrar la valorización)")
para("Si hubo adelanto, cada valorización lo amortiza con el mismo % de adelanto. El cobro neto al cliente es la valorización menos esa amortización.")
ejemplo("Valorización de S/ 821,250 con 20% de adelanto", ["Concepto", "Importe S/"], [
    ["Valorización de la semana", "821,250.00"],
    ["Amortización (20% × 821,250)", "164,250.00"],
    ["Cobro neto al cliente", "657,000.00"],
])
casos([
    ("Asigna contratista, fecha de inicio/entrega y duración a una partida hoja.", "El Estado pasa a 'En espera' si la fecha es futura, o se recalcula según el avance."),
    ("Pulsa 'Nueva valorización', escribe 25% de avance en una hoja y 'Guardar avances'.", "El % acumulado sube a 25% (barra verde), el saldo baja, y Estado/Prioridad se recalculan."),
    ("Intenta ingresar un avance que haga pasar el acumulado de 100%.", "El sistema lo bloquea y avisa."),
    ("Mira el panel de dilución y pulsa 'Registrar cobro'.", "El cobro neto (valorización − amortización) entra a la caja del proyecto y el saldo del adelanto baja."),
    ("Pulsa 'Resumen PDF'.", "Se genera el PDF de la valorización con el logo de AZUR."),
])
obs(3)

h2("7.5 Cronograma de cobros, Adicionales y Equipo")
botones([
    ["Agregar armada / Guardar cronograma", "Define cómo se cobrará (concepto, %, por avance o fecha)."],
    ["Registrar adicional/deductivo + Aprobar/Rechazar", "Gestiona cambios al contrato."],
    ["Asignar al equipo / Quitar", "Agrega personas con rol de obra (una persona puede tener varios roles)."],
])
obs(2)

h2("7.6 Mantenimiento (proyectos chicos)")
campos([
    ["Categoría", "Tipo de servicio (limpieza, luminarias…).", "Texto · obligatorio"],
    ["Descripción / Monto por visita", "Detalle y costo de cada visita.", "Texto / número"],
    ["Recurrencia", "Única/Semanal/Quincenal/Mensual/Trimestral/Semestral.", "Lista"],
    ["Inicio / N° de visitas", "Fecha de inicio y cuántas visitas generar.", "Fecha / número"],
    ["Avisar (días antes)", "7, 15 o 30 días antes de cada visita.", "Lista"],
])
botones([
    ["Generar cronograma", "Crea todas las visitas con sus fechas."],
    ["Ejecutado / Facturado", "Cambian el estado de cada servicio."],
])
para("El sistema genera una alerta automática cuando una visita se acerca o queda pendiente.")
obs(3)

h2("7.7 Campo (lo que llega del celular)")
pestanas([
    ["Asistencia", "Entradas/salidas con hora y enlace al mapa (GPS)."],
    ["Partes diarios", "RDO con actividades, avance, personal, observaciones."],
    ["Evidencias", "Galería de fotos por partida."],
    ["SST", "Charlas, observaciones e incidentes."],
])
obs(2)

h2("7.8 Liquidación")
tabla(["Línea", "Qué es"], [
    ["Contrato ajustado", "Contrato + adicionales − deductivos."],
    ["Valorizado / Cobrado / Por cobrar", "Avance facturado, cobrado y pendiente."],
    ["Costo presupuestado / Gastado real", "Presupuesto interno vs. egresos reales."],
    ["Margen vs. presupuesto / Utilidad real", "Resultado del proyecto y % de margen."],
    ["Saldo del adelanto", "Adelanto inicial − amortizado (baja hasta 0)."],
])
botones([
    ["PDF de liquidación", "Documento brandeado con el balance final."],
    ["Cerrar y liquidar obra", "Pone el proyecto en Liquidado, cierra la caja chica y devuelve el remanente a la caja central."],
])
obs(2)

# ───────────────────── 8. FINANZAS ─────────────────────
h1("8. Finanzas")
proposito("Trazabilidad total: solicitud → aprobación → pago → comprobante, además de CxC, CxP y cajas.")
h2("8.1 Solicitudes de pago — campos (desde la app o web)")
campos([
    ["Tipo", "Contratistas / Proveedores / Caja chica / Servicios / Honorarios.", "Lista · obligatorio"],
    ["Proyecto", "Proyecto al que se carga el gasto.", "Lista"],
    ["Partida presupuestal", "Partida del presupuesto.", "Texto/lista"],
    ["Beneficiario", "A quién se le paga.", "Texto"],
    ["Especialidad / Categoría-etapa", "Clasificación del gasto (ej. carpintería, MOD).", "Texto"],
    ["Monto (S/)", "Importe del gasto.", "Número · obligatorio"],
    ["Constancia", "Factura / Boleta / RHE.", "Lista"],
    ["Descripción", "Detalle del servicio/compra.", "Texto"],
    ["Cuenta bancaria", "Cuenta del beneficiario.", "Texto"],
    ["RUC / DNI · Razón social · N° comprobante", "Datos del beneficiario y del documento.", "Texto/números"],
])
para("La fecha de registro y el estado (Solicitada → Aprobada → Programada → Pagada → Conciliada) son automáticos.")
h2("8.2 Flujo de aprobación (3 niveles)")
botones([
    ["Aprobar (Jefe de Proyectos)", "Pasa a Aprobada y avisa al Administrador."],
    ["Rechazar / Devolver", "Devuelve con motivo al solicitante."],
    ["Programar (Administrador)", "Define banco de origen y fecha."],
    ["Pagar (Administrador)", "Registra método (transferencia/efectivo/Yape/Plin…), N° de operación y voucher adjunto."],
    ["Aprobar final (Gerencia)", "Solo si el monto supera el umbral configurado."],
    ["WhatsApp 📱 / Ojo 👁", "Envía el comprobante / muestra el detalle completo (17 campos)."],
])
pasos([
    ("Como Residente (en el celular) crea una solicitud y envíala.", "Queda Solicitada y avisa al Jefe."),
    ("Como Jefe de Proyectos, Aprobar.", "Pasa a Aprobada y avisa al Administrador."),
    ("Como Administrador, Programar y luego Pagar (método, N° operación, voucher).", "Pasa a Pagada; si supera el umbral, espera aprobación de Gerencia."),
    ("Como Gerencia, Aprobar final (si aplica).", "Queda Conciliada; si era caja chica, descuenta la caja; el solicitante recibe aviso."),
])
obs(3)
h2("8.3 Cuentas por cobrar (CxC)")
botones([
    ["Emitir factura (desde armada)", "Crea la factura de una armada pendiente."],
    ["Cobrar", "Registra un abono (parcial o total); actualiza el aging."],
    ["Factura manual", "Crea una factura suelta."],
])
para("Arriba ves el aging: Corriente / 1-30 / 31-60 / +60 días.")
obs(2)
h2("8.4 Cuentas por pagar (CxP) y Cajas")
botones([
    ["CxP", "Lista de obligaciones aprobadas/programadas pendientes de pago."],
    ["Abrir (caja)", "Abre la caja con su saldo e historial."],
    ["Registrar movimiento", "Reposición/abono/egreso/traslado/ajuste con método, N° operación y voucher."],
])
para("Cada caja avisa cuando se consume más del 80% de su tope.")
obs(2)

# ───────────────────── 9. ALMACÉN ─────────────────────
h1("9. Almacén (web)")
proposito("Controlar inventario y movimientos por proyecto.")
botones([
    ["Nuevo ítem", "Crea un ítem de inventario (código, nombre, unidad, stock, tipo)."],
    ["Registrar movimiento", "Ingreso (sin proyecto), Salida o Devolución (con proyecto). Actualiza el stock."],
])
para("Regla: el ingreso no lleva proyecto; las salidas y devoluciones sí.")
obs(2)

# ───────────────────── 10. REPORTES ─────────────────────
h1("10. Reportes")
proposito("Análisis financiero cruzado con filtros e indicadores.")
botones([
    ["Periodo (7/15/30 días, mes, histórico)", "Cambia el rango de análisis."],
    ["Filtro por línea / proyecto", "Acota los datos."],
    ["Exportar Excel", "Descarga el reporte brandeado según los filtros."],
])
para("Muestra ingresos vs egresos en el tiempo, resultados por línea, gasto por las 5 categorías y la tabla de proyectos con su salud.")
obs(2)

# ───────────────────── 11. ALERTAS ─────────────────────
h1("11. Alertas")
proposito("Centralizar los avisos del sistema (sobrecosto, hitos, cotizaciones, mantenimiento, etc.).")
botones([
    ["Filtro (todas/abiertas/resueltas)", "Cambia la vista."],
    ["Marcar resuelta", "Cierra la alerta."],
])
para("Muchas alertas se generan solas cada día (ver sección 13).")
obs(2)

# ───────────────────── 12. USUARIOS ─────────────────────
h1("12. Usuarios")
proposito("Administrar las cuentas y sus permisos.")
botones([
    ["Nuevo usuario", "Crea la cuenta (nombre, email, rol, teléfono, contraseña)."],
    ["Cambiar rol", "Modifica el rol desde la lista."],
    ["Activar / Desactivar", "Habilita o bloquea el acceso."],
    ["Contraseña", "Asigna una nueva contraseña al usuario."],
])
pasos([
    ("Crea un usuario y entra con él.", "Inicia sesión correctamente."),
    ("Entra como Residente y revisa que NO ve finanzas globales ni márgenes.", "Solo ve sus proyectos y su trabajo de campo."),
])
obs(2)

# ───────────────────── 13. APP DE OBRA (PWA) ─────────────────────
h1("13. App de obra (celular)")
proposito("Que el personal de obra registre todo desde el celular, incluso sin señal.")
para("Instala la app: abre la web en el celular y toca 'Instalar AZUR' (Android) o Compartir → Añadir a inicio (iPhone). Entra como Residente.")
h2("13.1 Inicio y asistencia")
botones([["Entrada / Salida (Check-in/out)", "Registra la asistencia con ubicación GPS y hora."]])
h2("13.2 Tareo de cuadrilla")
campos([["Proyecto / Fecha", "A qué obra y día.", "Lista/fecha"], ["Trabajador / horas / presente", "Cada miembro de la cuadrilla.", "Texto/número"]])
botones([["Agregar trabajador / Registrar tareo", "Añade filas y guarda la asistencia del personal."]])
h2("13.3 Parte diario (RDO)")
campos([
    ["Proyecto / Fecha", "Obra y día del parte.", "Lista/fecha"],
    ["Clima / Personal / Equipos / Materiales recibidos", "Condiciones del día.", "Texto/número"],
    ["Observaciones / Incidencias", "Notas del día.", "Texto"],
    ["Actividades (+ Actividad)", "Descripción, partida (opcional) y % de avance.", "Texto/número"],
])
botones([["Enviar parte del día", "Guarda el RDO y sus actividades; se consolida en el informe."]])
h2("13.4 Evidencias")
botones([["Tomar / elegir foto", "Abre la cámara o galería."], ["Subir evidencia", "Sube la foto con GPS y la vincula a la partida; aparece en la galería y en el proyecto."]])
h2("13.5 SST (seguridad)")
botones([["Charla 5 min / Observación / Incidente", "Registra cada uno (con foto opcional)."]])
h2("13.6 Almacén y modo sin conexión")
pasos([
    ("Registra una salida de material a un proyecto.", "Descuenta el stock."),
    ("Activa el modo avión, navega por pantallas ya vistas y llena un parte o solicitud.", "La app sigue funcionando; lo registrado se guarda y se sincroniza solo al volver la señal (aviso arriba)."),
])
h2("13.7 Notificaciones push")
pasos([
    ("Toca la campana y pulsa 'Activar notificaciones push'; acepta el permiso.", "Queda activado."),
    ("Haz que el Jefe apruebe una solicitud en otra ventana.", "Llega la notificación al celular y a la campana; al tocarla se marca leída."),
])
obs(3)

# ───────────────────── 14. AUTOMATIZACIÓN ─────────────────────
h1("14. Automatización diaria (alertas solas)")
proposito("Cada día (08:00 hora Perú) el sistema revisa y genera alertas/notificaciones sin que nadie las pida.")
tabla(["Revisa", "Qué hace"], [
    ["Cotizaciones", "Marca vencidas y avisa las por vencer."],
    ["Hitos", "Avisa los vencidos o próximos."],
    ["Armadas (cobros)", "Avisa las próximas a facturar."],
    ["Sobretiempo", "Proyectos con fecha de fin pasada y avance < 100%."],
    ["Salud financiera", "Gasto > cobrado / gasto > valorizado."],
    ["Cajas", "Cajas chicas sin movimiento por varios días."],
    ["Mantenimiento", "Servicios próximos según los días de aviso."],
    ["Jueves", "Recuerda emitir las valorizaciones de la semana."],
])
obs(2)

# ─────────────────── OBSERVACIONES GENERALES ───────────────────
doc.add_page_break()
h1("15. Observaciones generales y conclusiones")
para("Anota aquí comentarios globales, ideas de mejora o problemas que no entren en las secciones anteriores.")
t = doc.add_table(rows=1, cols=3); t.style = "Table Grid"
for j, x in enumerate(["Tema", "Prioridad (A/M/B)", "Detalle / sugerencia"]):
    c = t.rows[0].cells[j]; shade(c, "E20627"); cell_text(c, x, bold=True, white=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(12):
    rc = t.add_row()
    for j in range(3): cell_text(rc.cells[j], "", size=10)
doc.add_paragraph()
h2("Conformidad")
t = doc.add_table(rows=2, cols=2); t.style = "Table Grid"
cell_text(t.rows[0].cells[0], "Probado por (nombre y firma)", bold=True, size=10); cell_text(t.rows[0].cells[1], "", size=10)
cell_text(t.rows[1].cells[0], "Fecha", bold=True, size=10); cell_text(t.rows[1].cells[1], "", size=10)

page_number_footer()
# actualizar campos (TOC) al abrir en Word
upd = OxmlElement("w:updateFields"); upd.set(qn("w:val"), "true"); doc.settings.element.append(upd)

import shutil
repo_out = os.path.join("azur-erp", "GUIA_DE_PRUEBAS_AZUR.docx")
doc.save(repo_out)
print("Generado en repo:", os.path.abspath(repo_out))
try:
    shutil.copy(repo_out, "GUIA_DE_PRUEBAS_AZUR.docx")
    print("Copiado a la raíz.")
except PermissionError:
    alt = "GUIA_DE_PRUEBAS_AZUR_NUEVA.docx"
    shutil.copy(repo_out, alt)
    print("La raíz estaba bloqueada (¿Word abierto?). Guardé:", os.path.abspath(alt))
