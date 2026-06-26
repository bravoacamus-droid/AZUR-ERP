# -*- coding: utf-8 -*-
"""Guía de pruebas de las ÚLTIMAS actualizaciones AZUR ERP (.docx brandeado, español Perú)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, shutil

AZUR = RGBColor(0xE2, 0x06, 0x27)
AZUR2 = RGBColor(0xBE, 0x17, 0x23)
GRIS = RGBColor(0x66, 0x66, 0x66)
VERDE = RGBColor(0x0A, 0x7D, 0x33)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)
LOGO = next((p for p in ["public/logoazur.png", "logoazur.png", "../logoazur.png"] if os.path.exists(p)), None)

doc = Document()
st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10.5)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexcolor); tcPr.append(shd)

def cell_text(cell, text, bold=False, color=None, size=9.5, align=None, white=False):
    cell.text = ""; p = cell.paragraphs[0]
    if align: p.alignment = align
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
    if white: r.font.color.rgb = BLANCO
    elif color: r.font.color.rgb = color

def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs: r.font.color.rgb = AZUR; r.font.size = Pt(16)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs: r.font.color.rgb = AZUR2; r.font.size = Pt(13)
    return p

def para(text, color=GRIS, size=10):
    p = doc.add_paragraph(); r = p.add_run(text); r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return p

def proposito(text):
    p = doc.add_paragraph(); r = p.add_run("Qué se probó / para qué: "); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = AZUR2
    r2 = p.add_run(text); r2.font.size = Pt(10)
    return p

def donde(text):
    p = doc.add_paragraph(); r = p.add_run("Dónde: "); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = AZUR
    r2 = p.add_run(text); r2.font.size = Pt(10)
    return p

def pasos(items):
    for accion, esperado in items:
        p = doc.add_paragraph(style="List Number"); r = p.add_run(accion); r.font.size = Pt(10.5)
        if esperado:
            sp = doc.add_paragraph(); sp.paragraph_format.left_indent = Inches(0.5)
            rr = sp.add_run("✓ Debes ver: "); rr.bold = True; rr.font.size = Pt(9); rr.font.color.rgb = VERDE
            r2 = sp.add_run(esperado); r2.font.size = Pt(9); r2.font.color.rgb = GRIS

def casos(headers, rows):
    p = doc.add_paragraph(); r = p.add_run("Casos de prueba (hazlos y verifica el resultado):"); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = AZUR2
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    for j, x in enumerate(headers):
        c = t.rows[0].cells[j]; shade(c, "0A7D33"); cell_text(c, x, bold=True, white=True, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        rc = t.add_row()
        for j, val in enumerate(row): cell_text(rc.cells[j], val, size=9)
    doc.add_paragraph()

def nota(text):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run("💡 "); r.font.size = Pt(9.5)
    r2 = p.add_run(text); r2.font.size = Pt(9.5); r2.italic = True; r2.font.color.rgb = RGBColor(0x33,0x55,0x88)

def obs(filas=2):
    p = doc.add_paragraph(); r = p.add_run("📝 Observaciones de esta sección:"); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = AZUR2
    t = doc.add_table(rows=1, cols=4); t.style = "Table Grid"
    for j, x in enumerate(["Qué probaste", "¿Funcionó? Sí/No", "Severidad A/M/B", "Comentario"]):
        c = t.rows[0].cells[j]; shade(c, "E20627"); cell_text(c, x, bold=True, white=True, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(filas):
        rc = t.add_row()
        for j in range(4): cell_text(rc.cells[j], "", size=10)
    for col, w in zip(t.columns, [Inches(1.9), Inches(1.1), Inches(1.2), Inches(2.3)]):
        for c in col.cells: c.width = w
    doc.add_paragraph()

def add_toc():
    p = doc.add_paragraph(); run = p.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
    tt = OxmlElement("w:t"); tt.text = "Abre en Word y presiona F9 (clic derecho → Actualizar campos) para ver el índice con páginas."
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    for e in (f1, instr, f2, tt, f3): run._r.append(e)

def footer():
    fp = doc.sections[0].footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("AZUR Constructora e Inmobiliaria · Guía de Pruebas — Actualizaciones · Pág. ")
    r.font.size = Pt(8); r.font.color.rgb = GRIS
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE"); fp._p.append(fld)

# ───────────────────────── PORTADA ─────────────────────────
if LOGO:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(LOGO, width=Inches(2.2))
doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Guía de Pruebas — Actualizaciones del ERP"); r.bold = True; r.font.size = Pt(23); r.font.color.rgb = AZUR
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("AZUR Constructora e Inmobiliaria"); r.font.size = Pt(14); r.font.color.rgb = AZUR2
d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = d.add_run("Todas las observaciones levantadas por Juan y David, más las mejoras añadidas.\nPaso a paso, con el resultado esperado y espacio para anotar observaciones."); r.font.size = Pt(11); r.font.color.rgb = GRIS
doc.add_paragraph()
para("Versión de pruebas validada: el motor de cálculo (fórmulas de costo, plazos, valorización, "
     "adelanto y márgenes) fue verificado con 34 pruebas automáticas (34/34 OK) y el build de producción "
     "compila correctamente.", color=GRIS, size=9.5)
doc.add_page_break()

# ───────────────────────── CÓMO USAR ─────────────────────────
h1("Antes de empezar")
para("Sigue estos requisitos para que las pruebas reflejen la última versión:")
pasos([
    ("Abre la aplicación desde la URL desplegada (la del navegador) y presiona Ctrl+F5 para limpiar la caché.", "La versión más reciente cargada."),
    ("Ten a la mano usuarios con distintos roles: Gerencia, Jefe de proyectos y Presupuestos/Administrador.", "Algunas pruebas dependen del rol."),
    ("Trabaja sobre una cotización en estado Borrador o En negociación, y sobre un proyecto existente.", "Las ediciones solo se permiten en esos estados."),
    ("Anota lo que encuentres en la tabla de Observaciones al final de cada sección.", "Registro ordenado para el siguiente paquete de ajustes."),
])
nota("Severidad sugerida: A = bloqueante / impide trabajar · M = importante pero hay forma de continuar · B = detalle/estético.")
h2("Índice")
add_toc()
doc.add_page_break()

# ───────────────────────── 1. COTIZACIÓN ─────────────────────────
h1("1. Comercial — Cotización")

h2("1.1 Crear cotización (configuración inicial)")
proposito("Definir desde el arranque la moneda, el tipo de cambio, el plazo (ahora opcional) y el recomendado-por.")
donde("Menú Comercial → botón Nueva cotización.")
pasos([
    ("Elige Origen, Cliente, Línea, Tipo de cotización y Nombre del proyecto.", "Campos obligatorios marcados; permite crear cliente nuevo con el botón +."),
    ("En Moneda elige Dólares ($).", "Aparece el campo Tipo de cambio (S/ por $)."),
    ("Marca el check 'Definir plazo de ejecución'.", "Recién ahí aparece el número + tipo (días calendario / útiles / semanas / meses)."),
    ("Deja el check del plazo SIN marcar y crea la cotización.", "Se crea igual: el plazo es opcional y NO bloquea."),
    ("Cambia el Origen a 'Recomendación'.", "Aparece el campo 'Recomendado por'."),
])
obs()

h2("1.2 Moneda en dólares + tipo de cambio + equivalente en soles")
proposito("Cotizar en USD con el tipo de cambio que define el usuario, viendo el equivalente en soles (opcional).")
donde("Editor de la cotización → Bloque de totales → Configuración.")
pasos([
    ("Selecciona Moneda: Dólares y escribe el Tipo de cambio (ej. 3.75).", "Los totales se muestran en $ y aparece 'Equivalente en soles (T.C.): S/ …'."),
    ("Apaga el toggle 'Mostrar equivalente en soles'.", "Desaparece el equivalente del panel (y del PDF)."),
    ("Vuelve a Soles.", "Todo se muestra en S/."),
])
nota("Al aprobar la cotización en USD, el proyecto se crea con el contrato convertido a soles usando ese tipo de cambio (el proyecto y finanzas operan en soles).")
obs()

h2("1.3 Cuadro de costos: colores, unidades y costo unitario con fórmula")
proposito("Identificar partidas por color, elegir/escribir unidades, y calcular el costo unitario como en Excel.")
donde("Editor de la cotización → pestaña Presupuesto (cuadro de costos y margen).")
pasos([
    ("Observa el sombreado por nivel de las filas (partida fuerte, sub-partida medio, actividad claro).", "Colores por nivel, igual que el Last Planner."),
    ("En Unidad abre el desplegable y elige una; luego escribe una unidad libre que no esté en la lista.", "Acepta de la lista y también texto libre."),
    ("En una fila hoja (sin APU), en C. UNIT escribe (100+50)/2 y presiona Enter.", "Calcula 75; el P. Unit y Precio se recalculan."),
    ("Vuelve a hacer clic en esa celda de C. UNIT.", "Muestra la fórmula guardada (100+50)/2 para editarla."),
])
casos(["Escribes en C. UNIT", "Resultado esperado"], [
    ("=40/1.18  (quitar IGV)", "33.90"),
    ("4*50", "200"),
    ("=4*50+30", "230"),
    ("(100+50)/2", "75"),
    ("1200/12", "100"),
    ("40,5  (coma decimal)", "40.5"),
    ("abc  (texto no válido)", "no cambia / lo ignora"),
])
nota("P. Unit = C. Unit ÷ (1 − margen). Ej.: C.U. 100 y margen 30% → P. Unit 142.86.")
obs()

h2("1.4 Bloque de totales: qué ve el cliente + descuento (respuesta ágil)")
proposito("Controlar qué conceptos ve el cliente y aplicar descuento; los cambios deben sentirse inmediatos.")
donde("Editor de la cotización → Bloque de totales → Configuración.")
pasos([
    ("Apaga/enciende 'Mostrar GG', 'GA', 'Utilidad', 'IGV'.", "El panel se actualiza al instante (sin recargar). Eso controla lo que sale en PDF/Excel."),
    ("Agrega un descuento comercial y cambia el %.", "El TOTAL CON DESCUENTO se recalcula al momento."),
])
obs()

h2("1.5 PDF de la cotización")
proposito("Verificar el PDF brandeado con moneda, forma de pago y firma del responsable.")
donde("Editor de la cotización → Acciones / botón de PDF.")
pasos([
    ("Descarga el PDF.", "Sale con la moneda elegida; si es USD, muestra T.C. y equivalente en soles."),
    ("Revisa la sección FORMA DE PAGO.", "Lista las armadas (adelanto, valorización, etc.) con su % y monto."),
    ("Revisa el pie del documento.", "Bloque de firma con el nombre y cargo del responsable."),
])
obs()

h2("1.6 Excel de la cotización (celdas con fórmula)")
proposito("Que al editar en Excel los totales se recalculen (celdas referenciadas), no valores fijos.")
donde("Editor de la cotización → Descargar Excel.")
pasos([
    ("Abre el Excel y haz clic en la celda SUBTOTAL (columna F) de una partida.", "En la barra de fórmulas verás =D*E (cantidad × costo unitario), no un número fijo."),
    ("Cambia la CANT (columna D) de una fila.", "El SUBTOTAL y el SUBTOTAL C/M se recalculan solos."),
    ("Si en la app escribiste el costo como fórmula (ej. (100+50)/2), haz clic en C. UNITARIO (columna E).", "La barra muestra =(100+50)/2 (la fórmula que escribiste); si no, muestra el número."),
])
nota("La columna C. UNITARIO es la ENTRADA: muestra el valor (75) o la fórmula que escribiste. Las fórmulas referenciadas de David están en SUBTOTAL, P. UNIT, MARGEN y SUBTOTAL C/M.")
obs()

# ───────────────────────── 2. LAST PLANNER ─────────────────────────
doc.add_page_break()
h1("2. Proyecto — Last Planner")
donde("Menú Proyectos → abre un proyecto → pestaña Last Planner.")

h2("2.1 Visual: numeración y color por nivel")
pasos([
    ("Observa la primera columna (Ítem) y el color de las filas.", "Numeración automática 1.0 / 1.1 / 1.1.1 y sombreado por nivel (incluidas partidas hoja como 2.0/3.0)."),
    ("Edita la Unidad (lista + texto libre) y el C. UNIT (acepta fórmula en filas hoja sin APU).", "Igual que en la cotización."),
])
obs()

h2("2.2 Plazos: calendario + Recalcular fechas")
proposito("Calcular la fecha de entrega a partir de inicio + duración, contando solo los días del calendario elegido.")
pasos([
    ("En la barra superior elige el calendario (ej. Solo sábado y domingo).", "Queda seleccionado."),
    ("En una fila pon Fecha de inicio y Duración (ej. 5).", "Quedan registradas."),
    ("Pulsa 'Recalcular fechas'.", "TODAS las entregas se recalculan de golpe (con aviso); cuentan solo los días del calendario."),
    ("Abre una fecha y haz clic afuera SIN cambiar nada.", "Se queda igual: ya NO se corre un día (bug de zona horaria corregido)."),
])
casos(["Calendario / Inicio / Duración", "Entrega esperada"], [
    ("Lunes a viernes · un lunes · 5", "el viernes de esa semana"),
    ("Lunes a viernes · un lunes · 6", "el lunes siguiente (salta sábado y domingo)"),
    ("Lunes a sábado · un lunes · 6", "el sábado de esa semana"),
    ("Solo sábado y domingo · un sábado · 4", "dos fines de semana después"),
])
nota("La fecha de INICIO no se recalcula: es el punto de partida que defines tú. Se calculan Entrega y Duración entre sí.")
obs()

h2("2.3 Valorización: % y monto en paralelo + acumulado + fila TOTALES")
proposito("Llenar el avance por % o por monto (lo que falte se calcula solo) y ver los acumulados.")
pasos([
    ("Crea una valorización (botón Nueva valorización) si hace falta.", "Aparecen las columnas '% sem' y 'S/ sem' del periodo."),
    ("En una fila escribe el %; en otra escribe el monto S/.", "Si pones %, calcula el monto; si pones el monto, calcula el %."),
    ("Revisa las columnas 'Valorizado' (acumulado) y 'Saldo'.", "Acumulado en soles y saldo por partida."),
    ("Baja al pie de la tabla (fila TOTALES).", "Muestra el total contractual, valorizado acumulado, saldo y el TOTAL de cada valorización (como el Excel)."),
])
casos(["En una partida de Total 2,000", "Resultado"], [
    ("Escribes 10 en %", "monto = S/ 200"),
    ("Escribes 500 en S/", "% = 25%"),
])
obs()

h2("2.4 Base de valorización: Costo o Precio (con margen)")
proposito("Elegir si la valorización/cobro se calcula sobre el costo del itemizado o sobre el precio (con margen).")
donde("Proyecto → Resumen → Configuración rápida → 'Base de valorización / cobro'.")
pasos([
    ("Cambia a 'Sobre precio (con margen)'.", "En el Last Planner, los montos de valorización/saldo/totales suben al precio de venta; la tarjeta muestra el badge 'Base: Precio'."),
    ("Vuelve a 'Sobre costo'.", "Regresan a los valores del itemizado."),
])
nota("Factor = contrato ÷ costo directo. Al 100% de avance, en 'Precio' el valorizado total = contrato. Por defecto está en 'Costo' (no cambia nada de lo actual). Pendiente: que Juan defina cuál usarán.")
obs()

h2("2.5 Aprobaciones de cantidad / costo (Presupuestos)")
proposito("Que un cambio de monto del Jefe de proyectos pase por aprobación de Presupuestos.")
pasos([
    ("Entra como Jefe de proyectos (en un proyecto con itemizado heredado de la cotización).", "—"),
    ("Cambia una Cantidad o un C. UNIT de una partida.", "Aviso: 'enviado a aprobación de Presupuestos'; el valor vuelve al original (no se aplica todavía)."),
    ("Entra como Presupuestos o Gerencia → Resumen → 'Solicitudes de cambio pendientes' → Aprobar.", "El cambio queda aplicado."),
])
nota("Si el proyecto usa 'Itemizado propio', el Jefe edita libre (no hay un comercial aprobado que proteger).")
obs()

h2("2.6 Corregir / reabrir una valorización ya cobrada")
proposito("Corregir una valorización pasada con el control adecuado según el rol.")
pasos([
    ("Como GERENCIA: sobre una valorización anterior pulsa 'Reabrir N°X'.", "Se reabre al instante (con confirmación) — Gerencia no se pide permiso a sí misma."),
    ("Como Jefe de proyectos / Presupuestos: pulsa 'Solicitar corrección N°X'.", "Envía la solicitud a Gerencia, que la aprueba en Resumen → Solicitudes."),
    ("Edita el avance y pulsa Guardar avances.", "Al guardar, la valorización se vuelve a bloquear."),
])
obs()

h2("2.7 Itemizado propio + comparativo Comercial vs Proyecto")
pasos([
    ("Resumen → 'Itemizado de Proyectos' → activa 'Itemizado propio' y/o 'Vaciar itemizado heredado'.", "Permite armar una estructura distinta a la cotización."),
    ("Arma tus partidas en el Last Planner y vuelve a Resumen.", "La tarjeta 'Comparativo Comercial vs Proyecto' compara por totales y margen aunque las categorías no coincidan."),
])
obs()

h2("2.8 Presupuesto por tipo de gasto")
pasos([
    ("Resumen → 'Presupuesto por tipo de gasto': escribe el proyectado por tipo (contratistas, proveedores, caja chica, servicios, honorarios).", "La suma debe cuadrar con el costo (✓)."),
    ("Revisa la columna Real y el Gap.", "Real = solicitudes pagadas/conciliadas por tipo; Gap = proyectado − real."),
])
obs()

h2("2.9 Adelantos (contrato + adicional / extraordinario)")
proposito("Registrar adelantos extra que se diluyen junto con el del contrato.")
pasos([
    ("Resumen → 'Adelantos': revisa el adelanto del contrato (%).", "Muestra el monto del adelanto contractual."),
    ("Registra un adelanto adicional/extraordinario (concepto + monto).", "El total de adelanto sube; en la valorización, la amortización por periodo aumenta y el saldo del adelanto baja hasta 0 al cerrar la obra."),
])
casos(["Contrato 100,000 · adelanto 20% · valorización del periodo", "Amortización / Cobro neto"], [
    ("10,000", "Amortización 2,000 · Cobro neto 8,000"),
])
obs()

# ───────────────────────── 3. INFORME ─────────────────────────
doc.add_page_break()
h1("3. Informe / PDF de la valorización")
donde("Last Planner → tarjeta de la valorización activa → botón Resumen PDF.")
pasos([
    ("Genera el Resumen PDF.", "Resumen ejecutivo: Monto del contrato → (−) Adelanto recibido → Periodo (valorización, amortización, cobro neto) → Acumulado (valorizado, adelanto amortizado, saldo del adelanto, saldo por valorizar)."),
    ("Revisa 'Valorizaciones acumuladas'.", "Lista Val 1, 2, 3… con sus montos sumando al acumulado."),
    ("Revisa 'Detalle por partida'.", "Ítem, unidad, contractual, % periodo, valorizado, % acumulado, valorizado acumulado, saldo + fila de TOTALES."),
    ("Revisa el pie.", "Firmas: Elaborado por (Jefe de Proyectos) y Aprobado por (Gerencia)."),
])
obs()

# ───────────────────────── 4. FINANZAS ─────────────────────────
doc.add_page_break()
h1("4. Finanzas")

h2("4.1 Crear requerimiento de pago desde la web")
donde("Menú Finanzas → pestaña Solicitudes de pago.")
pasos([
    ("Pulsa 'Nueva solicitud de pago' y llena (tipo, proyecto, beneficiario, monto, RUC/CCI, etc.).", "Se crea como 'Solicitada' y entra al flujo de aprobación."),
])
obs()

h2("4.2 Estados de pago y roles")
pasos([
    ("Lee la leyenda del flujo en la parte superior de Solicitudes.", "Solicitada (campo) → Aprobada (Jefe) → Programada (Administración) → Pagada (Gerencia) → Conciliada."),
    ("Verifica los botones según el rol.", "Solo Gerencia ve 'Pagar'; Administración solo 'Programar'."),
])
obs()

h2("4.3 Estado financiero por proyecto")
donde("Finanzas → pestaña Cajas (parte superior).")
pasos([
    ("Revisa la tabla 'Estado financiero por proyecto'.", "Contrato / Cobrado / Gastado / Por cobrar / Disponible por cada proyecto."),
])
obs()

h2("4.4 Caja chica por residente — control semanal")
donde("Finanzas → pestaña Cajas.")
pasos([
    ("Pulsa 'Nueva caja chica' y asigna proyecto + responsable + asignación semanal.", "Se crea una caja chica por residente."),
    ("Abre la caja y registra reposiciones (entregado) y egresos (gastado) con fechas.", "—"),
    ("Revisa la tabla 'Control semanal'.", "Entregado vs gastado por semana, saldo acumulado y alerta de sobregasto / exceso de entrega."),
])
obs()

h2("4.5 Datos bancarios y detracción (proveedor / cliente)")
donde("Catálogos → Contrapartes (editar) y Clientes → Maestro (editar).")
pasos([
    ("Edita un proveedor: revisa RUC, CCI y 'Cuenta de detracción'.", "Campos disponibles."),
    ("En 'Cuentas bancarias adicionales' agrega más de una cuenta (banco, moneda, cuenta, CCI, detracción).", "Permite varias cuentas por proveedor/cliente."),
    ("Edita un cliente: revisa banco, cuenta, CCI y cuenta de detracción.", "Disponibles también para el cliente."),
])
nota("Para agregar cuentas adicionales, primero guarda el proveedor/cliente (por la relación).")
obs()

# ───────────────────────── 5. REPORTES ─────────────────────────
doc.add_page_break()
h1("5. Reportes")
donde("Menú Reportes (filtros de periodo / proyecto / línea).")
pasos([
    ("Revisa 'Proyectado vs Real por tipo de gasto'.", "Gráfico de barras + tabla con el Gap por categoría."),
    ("Revisa la Curva S.", "Ejes rotulados: X 'Semanas (cronograma)', Y '% avance acumulado', con leyenda Planificado/Ejecutado."),
    ("Revisa 'Ingresos vs. egresos en el tiempo'.", "Ejes rotulados: X 'Fecha', Y 'Monto (S/)', con leyenda."),
])
obs()

# ───────────────────────── 6. TRANSVERSALES ─────────────────────────
doc.add_page_break()
h1("6. Detalles transversales (correcciones)")
pasos([
    ("Abre pantallas con tarjetas KPI (Inicio, Finanzas, Reportes, Proyecto) y prueba a maximizar / usar doble pantalla / zoom.", "Los íconos ya NO se solapan con los montos (se ajustan y, si el número es muy grande, se recorta mostrándolo completo al pasar el mouse)."),
    ("Revisa cualquier fecha del sistema (al ver, abrir o cerrar).", "Las fechas ya NO se corren un día (corrección de zona horaria, global)."),
])
obs()

# ───────────────────────── 7. REGISTRO GENERAL ─────────────────────────
doc.add_page_break()
h1("7. Registro general de observaciones")
para("Usa esta tabla para un resumen final de los hallazgos prioritarios.")
obs(filas=12)

para("Validación interna: motor de cálculo verificado con 34 pruebas automáticas (34/34 OK) y build de "
     "producción compilando sin errores. Esta guía cubre las observaciones de Juan y David más las mejoras añadidas.",
     color=GRIS, size=9)

footer()

# ───────────────────────── GUARDAR ─────────────────────────
NOMBRE = "GUIA_PRUEBAS_ACTUALIZACIONES_AZUR.docx"
destinos = [os.path.join("..", NOMBRE), NOMBRE]  # raíz del proyecto + carpeta del repo
guardados = []
for dest in destinos:
    try:
        doc.save(dest); guardados.append(os.path.abspath(dest))
    except PermissionError:
        print("AVISO: cierra el archivo si está abierto en Word ->", dest)
print("OK guardado en:")
for g in guardados: print("  -", g)
