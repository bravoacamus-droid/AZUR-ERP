# -*- coding: utf-8 -*-
"""Guía de pruebas AZUR ERP (.docx brandeado) — misma guía que el artefacto web."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

AZUR = RGBColor(0xE2, 0x06, 0x27)
AZUR2 = RGBColor(0xBE, 0x17, 0x23)
GRIS = RGBColor(0x66, 0x66, 0x66)
TINTA = RGBColor(0x1A, 0x15, 0x16)
VERDE = RGBColor(0x0A, 0x7D, 0x33)
WARN = RGBColor(0x9A, 0x6A, 0x00)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)
LOGO = next((p for p in ["public/logoazur.png", "logoazur.png"] if os.path.exists(p)), None)

doc = Document()
for s in doc.sections:
    s.top_margin = Inches(0.7); s.bottom_margin = Inches(0.7)
    s.left_margin = Inches(0.8); s.right_margin = Inches(0.8)
st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10.5); st.font.color.rgb = TINTA

def shade(el, hexcolor):
    tcPr = el.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexcolor); tcPr.append(shd)

def hrule(color="E20627", size=12):
    p = doc.add_paragraph(); pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), color)
    pbdr.append(bottom); pPr.append(pbdr); p.paragraph_format.space_after = Pt(2)
    return p

# ── Cabecera ──────────────────────────────────────────────────────────
tbl = doc.add_table(rows=1, cols=2); tbl.autofit = True
c0, c1 = tbl.rows[0].cells
if LOGO:
    c0.paragraphs[0].add_run().add_picture(LOGO, width=Inches(0.95))
c0.width = Inches(1.1)
p = c1.paragraphs[0]
r = p.add_run("AZUR"); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = AZUR
r2 = p.add_run("  CONSTRUCTORA E INMOBILIARIA"); r2.font.size = Pt(8); r2.font.color.rgb = GRIS
sub = c1.add_paragraph(); rs = sub.add_run("Guía de pruebas · Actualizaciones (Comercial · Roles · Firmas · Proyecto)")
rs.font.size = Pt(10.5); rs.bold = True; rs.font.color.rgb = TINTA
hrule()

intro = doc.add_paragraph()
ri = intro.add_run("Cada punto indica qué cambió y los pasos exactos para probarlo, con el resultado esperado. "
                   "Todo está desplegado; si algo no aparece, recarga con Ctrl+F5. Marca con un check cada prueba superada.")
ri.font.size = Pt(10); ri.font.color.rgb = GRIS
doc.add_paragraph()

# ── Helpers de contenido ──────────────────────────────────────────────
def seccion(num, titulo):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
    rn = p.add_run(num + "  "); rn.bold = True; rn.font.color.rgb = AZUR; rn.font.size = Pt(13)
    rt = p.add_run(titulo); rt.bold = True; rt.font.size = Pt(14); rt.font.color.rgb = TINTA
    hrule(size=10)

def item(tag, titulo, what, steps, esperado):
    # Título con casilla
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(0)
    rc = p.add_run("☐  "); rc.font.size = Pt(12)
    rt = p.add_run(titulo); rt.bold = True; rt.font.size = Pt(11.5); rt.font.color.rgb = TINTA
    rg = p.add_run("    " + tag.upper()); rg.font.size = Pt(7.5); rg.font.color.rgb = GRIS
    # Qué cambió
    pw = doc.add_paragraph(); pw.paragraph_format.left_indent = Inches(0.28); pw.paragraph_format.space_after = Pt(2)
    rw = pw.add_run(what); rw.italic = True; rw.font.size = Pt(9.5); rw.font.color.rgb = GRIS
    # Pasos
    for s in steps:
        ps = doc.add_paragraph(style="List Number"); ps.paragraph_format.left_indent = Inches(0.55)
        rs = ps.add_run(s); rs.font.size = Pt(10)
    # Esperado
    pe = doc.add_paragraph(); pe.paragraph_format.left_indent = Inches(0.28); pe.paragraph_format.space_before = Pt(1)
    re1 = pe.add_run("Esperado:  "); re1.bold = True; re1.font.size = Pt(9.5); re1.font.color.rgb = VERDE
    re2 = pe.add_run(esperado); re2.font.size = Pt(9.5); re2.font.color.rgb = TINTA

# ── Contenido (igual que la guía web) ─────────────────────────────────
DATA = [
 ("01", "Comercial — Cotizaciones", [
  ("Datos generales", "Cambiar cliente, asunto y ubicación",
   "Antes no se podía cambiar el cliente de una copia; ahora sí.",
   ["Abre una cotización (por ejemplo una copia).",
    "Ve a la pestaña Condiciones y pago.",
    "En el bloque Datos generales cambia el Cliente en el desplegable.",
    "Edita también Asunto y Ubicación si quieres."],
   "El nombre del cliente cambia al instante en la cabecera y queda guardado."),
  ("Itemizado", "Reordenar partidas (subir / bajar)",
   "Mueve una partida entre sus hermanas; el código se renumera solo.",
   ["En el Cuadro de costos, usa las flechas ↑/↓ a la izquierda de cada fila.",
    "Sube la 1.3 → pasa a 1.2, 1.1…",
    "Sube una partida 2.0 → pasa a 1.0."],
   "Se reacomoda al instante y el código (1.1, 1.2, 2.0…) se recalcula solo."),
  ("Itemizado", "Título de partida multilínea",
   "Agrupar servicios en un mismo bloque con saltos de línea.",
   ["Haz clic en el título de una partida.",
    "Presiona Enter para agregar una línea dentro del mismo bloque.",
    "Genera el PDF."],
   "El PDF respeta los saltos de línea del título."),
  ("Itemizado", "Navegación tipo Excel (flechas, Enter y Tab)",
   "Moverse entre celdas sin hacer clic afuera, como en Excel.",
   ["Párate en una celda (título, cantidad, costo…).",
    "↑/↓: salta a la misma columna de la fila anterior/siguiente.",
    "Tab: avanza a la siguiente celda editable (izquierda→derecha) y salta de fila al final; Shift+Tab retrocede.",
    "Enter: baja una fila (dentro del título multilínea, agrega salto de línea)."],
   "Tab se salta los botones de acción y se mueve solo entre celdas editables; en el borde deja el Tab normal para salir de la tabla."),
  ("Importar", "Importar itemizado desde Excel",
   "Plantilla brandeada AZUR + parser que ubica la cabecera.",
   ["Botón Importar Excel (junto a “Agregar partida”).",
    "Clic en Plantilla: descarga un .xlsx brandeado (barra roja AZUR, cabecera, ejemplos).",
    "Complétala y súbela, o copia/pega las filas.",
    "Revisa la vista previa y pulsa Importar."],
   "Se crean las partidas con jerarquía por Nivel (1=partida, 2=sub partida). Importa valores, no fórmulas."),
  ("Totales", "GG / GA / Utilidad en 0% y ocultos",
   "La nueva cotización nace limpia; solo IGV visible.",
   ["Crea una nueva cotización.",
    "Mira el bloque de totales."],
   "GG, GA y utilidad nacen en 0% y no visibles; solo el IGV. Tú los activas y editas si los necesitas."),
  ("Condiciones", "Viñetas automáticas",
   "Bullets ordenados sin poner asteriscos a mano.",
   ["Condiciones y pago → campos Condiciones / Servicios incluidos / omitidos.",
    "Escribe: arranca con • ; cada Enter agrega otra viñeta.",
    "Genera el PDF."],
   "Las viñetas salen ordenadas en el PDF."),
  ("Lista", "Pestaña “Plantillas” separada",
   "La lista ya no mezcla plantillas ni borradores.",
   ["En Comercial, alterna las pestañas Cotizaciones y Plantillas."],
   "Las plantillas viven en su pestaña; la lista de cotizaciones queda limpia."),
  ("Bug", "Plazo con decimales",
   "Antes 3.5 no se guardaba (la columna era entera).",
   ["Condiciones y pago → Plazo de ejecución → escribe 3.5 semanas.",
    "Sal del campo y vuelve a entrar."],
   "Ahora sí se guarda 3.5."),
  ("Bug", "Condiciones generales en el PDF",
   "Antes no salían en el PDF.",
   ["Llena “Condiciones generales”.",
    "Genera el PDF de cliente."],
   "Aparece la sección CONDICIONES GENERALES en el PDF."),
  ("Bug", "Menú de 3 puntos",
   "No abría (un stopPropagation cortaba el clic).",
   ["En la lista de Comercial, clic en los 3 puntos de una fila."],
   "El menú abre: Descargar PDF · Copiar link · Enviar a revisión · Duplicar · Solicitar/Eliminar."),
  ("Flujo (notifica)", "Revisión / validación",
   "Reemplaza el “mándame el archivo por WhatsApp”.",
   ["Dentro de la cotización o desde el menú: Enviar a revisión.",
    "Presupuestos/Gerencia reciben notificación (campana + push) con el link.",
    "El revisor abre y usa Aprobar revisión u Observar (con nota).",
    "Se notifica de vuelta a quien la envió."],
   "Badges En revisión / Aprobada / Observada; la observación queda visible para corregir."),
  ("Aprobación (notifica)", "Eliminar con aprobación",
   "Solo Gerencia elimina; el resto solicita.",
   ["Si NO eres Gerencia: menú → Solicitar eliminación (notifica a Gerencia).",
    "Como Gerencia: abre la cotización → banner con Aprobar eliminación / Cancelar."],
   "No se puede eliminar una cotización Aceptada (ya generó proyecto)."),
 ]),
 ("02", "Roles y usuarios", [
  ("Seguridad", "Roles personalizados (por módulo)",
   "Crea roles a medida: por cada módulo, Sin acceso / Solo ver / Editar.",
   ["Usuarios → Nuevo rol.",
    "Nómbralo y marca por módulo: Sin acceso / Solo ver / Editar.",
    "Asígnalo a un usuario (en su fila o al crearlo)."],
   "Un “Solo ver” no puede crear/editar: bloqueado en el servidor y con los botones ocultos."),
  ("Usuarios", "Editar datos del usuario",
   "Antes no se podía cambiar nombre/email/teléfono.",
   ["Usuarios → Editar en la fila.",
    "Cambia nombre, email o teléfono."],
   "Se guardan los cambios (el email se actualiza también en el acceso)."),
  ("Bug", "Rol personalizado al crear usuario",
   "Faltaba el selector en el alta.",
   ["Usuarios → Nuevo usuario.",
    "En Rol personalizado (opcional) elige uno de la lista."],
   "Si no hay roles aún, dice “crea uno con Nuevo rol”. También se asigna después desde la lista."),
  ("Seguridad", "Botones ocultos en “Solo ver”",
   "Coherencia visual con el permiso.",
   ["Con un usuario “Solo ver”, entra a Finanzas / Almacén / Catálogos / Clientes."],
   "No aparecen los botones Nuevo/Editar/Borrar (además el servidor los bloquea)."),
 ]),
 ("03", "Firmas", [
  ("Perfil", "Firma por usuario",
   "Cada usuario carga su firma PNG.",
   ["Mi perfil → bloque Mi firma → sube un PNG (ideal sin fondo).",
    "O como admin: Usuarios → Editar → Firma."],
   "Se ve la previsualización; puedes quitarla y reemplazarla."),
  ("Documentos", "Firmantes configurables por documento",
   "Elige quién firma cada PDF con checkboxes.",
   ["Cotización: Condiciones y pago → Firmas del documento → marca firmantes.",
    "Proyecto: Resumen → Firmas del documento (aplica a valorización y liquidación).",
    "Cada usuario muestra su rol y si tiene Firma ✓ o Sin firma.",
    "Genera el PDF."],
   "La firma se incrusta si el usuario la tiene cargada; si no, queda el espacio para firmar. Verificado por render."),
  ("Automático", "Firma del gerente único",
   "Como hay un solo gerente, va automática.",
   ["Sin configurar firmantes, genera una valorización."],
   "En “Aprobado por · Gerencia” sale el nombre y la firma del gerente (si la tiene cargada)."),
  ("Visual", "Más aire en las firmas",
   "Ya no quedan ajustadas al texto de arriba.",
   ["Genera una valorización y mira el bloque de firmas."],
   "Hay más espacio arriba de las líneas de firma."),
 ]),
 ("04", "Cuentas de empresa y PDFs", [
  ("Catálogos", "Cuentas reales + dónde se muestran",
   "Cuentas AZUR reales; eliges el documento por cuenta.",
   ["Catálogos → Medios de pago.",
    "Cada cuenta tiene toggles Mostrar en: Cotización / Valorización / Liquidación.",
    "Genera cada PDF."],
   "Cada cuenta aparece solo en los documentos marcados (BBVA soles + Interbank soles/dólares reales)."),
  ("PDF", "Medios de pago en valorización/liquidación",
   "El cliente ve dónde depositar.",
   ["Genera una valorización o liquidación."],
   "Aparece el bloque “Medios de pago — depositar a nombre de…”."),
  ("PDF", "Cabecera fija + tabla que se repite + firma nunca sola",
   "Documentos multipágina bien armados.",
   ["Genera una valorización con muchas partidas (varias páginas)."],
   "La cabecera AZUR y el encabezado de columnas se repiten en cada hoja; la firma nunca queda sola. Verificado por render."),
 ]),
 ("05", "Proyecto / Dashboard", [
  ("Bug clave", "Comparativo Comercial vs Proyecto en soles",
   "Antes mezclaba dólares y soles (desviación falsa).",
   ["Abre un proyecto en dólares (ej. Pelícanos) → Resumen → Comparativo Comercial vs Proyecto."],
   "Compara en soles: márgenes iguales (~42.7% vs 42.7%) y desviación ~0. Badge USD → S/ · T.C. 3.500 + tips (?) que explican cada número."),
  ("Fluidez", "Optimismo: agregar/borrar/restaurar sin recargar",
   "La UI responde al instante.",
   ["Agrega o borra una partida; restaura una versión guardada."],
   "No recarga la página; el spinner de restauración es por-versión y termina con “Restaurado ✓”."),
 ]),
]

for num, titulo, items in DATA:
    seccion(num, titulo)
    for it in items:
        item(*it)

doc.add_paragraph()
hrule()
nota = doc.add_paragraph()
rn = nota.add_run("Notas: ")
rn.bold = True; rn.font.color.rgb = WARN; rn.font.size = Pt(9.5)
rn2 = nota.add_run("Las notificaciones push llegan a quien las tenga activadas en su dispositivo (la campana in-app siempre funciona). "
                   "Las firmas de valorización/liquidación se configuran por proyecto (una vez), no por cada valorización.")
rn2.font.size = Pt(9.5); rn2.font.color.rgb = GRIS

foot = doc.add_paragraph(); foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = foot.add_run("AZUR Constructora e Inmobiliaria · Guía de pruebas de actualizaciones")
rf.font.size = Pt(8); rf.font.color.rgb = GRIS

OUT = []
for base in [".", ".."]:
    try:
        path = os.path.join(base, "GUIA_PRUEBAS_AZUR.docx")
        doc.save(path); OUT.append(os.path.abspath(path))
    except Exception as e:
        print("No se pudo guardar en", base, ":", e)
print("Guardado en:")
for p in OUT: print(" -", p)
